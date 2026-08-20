"""Xuất file DOCX theo chuẩn Thể thức Văn bản Hành chính Công vụ Việt Nam.

Tuân thủ:
- Nghị định 30/2020/NĐ-CP về công tác văn thư nhà nước (Quốc hiệu, Tiêu ngữ, Căn lề A4, Phông chữ Times New Roman 13pt, Canh đều 2 bên Justified, Thụt đầu dòng 1.27cm).
- Mẫu E-HSMT chuẩn ban hành kèm theo Thông tư 22/2024/TT-BKHĐT và Nghị định 214/2025/NĐ-CP.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Mm, Pt, RGBColor

from autotender.hitl.store import HitlStore
from autotender.models.generator import CHAPTER_TITLES
from autotender.schemas import HSMTDocument, HSMTSection

_CHAPTER_RANK = {chapter: i for i, chapter in enumerate(CHAPTER_TITLES)}


def _ordered_sections(sections: list[HSMTSection]) -> list[HSMTSection]:
    """Sắp theo thứ tự chương CHÍNH THỨC (I→VIII)."""
    return sorted(sections, key=lambda s: _CHAPTER_RANK.get(s.section_id.split(".")[0], len(_CHAPTER_RANK)))


def _set_cell_background(cell, fill_hex: str) -> None:
    """Tô màu nền cho ô trong bảng (Table Cell Shading)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def _set_cell_margins(cell, top: int = 120, bottom: int = 120, left: int = 150, right: int = 150) -> None:
    """Đặt padding cho ô trong bảng (đơn vị dxa, 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _add_page_number_to_footer(footer) -> None:
    """Thêm trường số trang động 'Trang X / Y' vào Footer."""
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    fldSimple1 = OxmlElement("w:fldSimple")
    fldSimple1.set(qn("w:instr"), "PAGE")
    p._p.append(fldSimple1)

    run2 = p.add_run(" / ")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(100, 100, 100)

    fldSimple2 = OxmlElement("w:fldSimple")
    fldSimple2.set(qn("w:instr"), "NUMPAGES")
    p._p.append(fldSimple2)


def export_docx(doc: HSMTDocument, store: HitlStore, output_path: str | Path) -> Path:
    """Xuất file DOCX chuẩn Thể thức Văn bản Hành chính (NĐ 30/2020 & Thông tư 22/2024)."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    # 1. CẤU HÌNH KHỔ TRANG A4 & LỀ THEO NGHỊ ĐỊNH 30/2020/NĐ-CP (Mục 8 SPEC)
    # Lề trên: 20mm, lề dưới: 20mm, lề trái: 30mm (đóng gáy), lề phải: 20mm
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)

    # 2. CẤU HÌNH PHÔNG CHỮ & STYLES CHUẨN TIMES NEW ROMAN
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(13)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)

    # Thêm số trang vào Footer
    _add_page_number_to_footer(section.footer)

    # =========================================================================
    # PHẦN I: TRANG BÌA HỒ SƠ MỜI THẦU (COVER PAGE CHUẨN CÔNG VỤ)
    # =========================================================================

    # Header 2 cột: Cơ quan ban hành (trái) vs Quốc hiệu Tiêu ngữ (phải)
    header_table = document.add_table(rows=2, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False

    col_widths = [Mm(75), Mm(85)]
    for row in header_table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    # Cột trái: Đơn vị chủ quản & Bên mời thầu
    c_left_top = header_table.cell(0, 0)
    p_cq = c_left_top.paragraphs[0]
    p_cq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cq.paragraph_format.space_after = Pt(2)
    r_cq = p_cq.add_run(f"{(doc.package.investor or 'CƠ QUAN CHỦ ĐẦU TƯ').upper()}")
    r_cq.font.size = Pt(11)
    r_cq.font.bold = True

    c_left_bot = header_table.cell(1, 0)
    p_bmt = c_left_bot.paragraphs[0]
    p_bmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bmt.paragraph_format.space_after = Pt(0)
    r_bmt = p_bmt.add_run("BÊN MỜI THẦU\n───────────")
    r_bmt.font.size = Pt(11)
    r_bmt.font.bold = True

    # Cột phải: Quốc hiệu và Tiêu ngữ
    c_right_top = header_table.cell(0, 1)
    p_qh = c_right_top.paragraphs[0]
    p_qh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qh.paragraph_format.space_after = Pt(2)
    r_qh = p_qh.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    r_qh.font.size = Pt(11)
    r_qh.font.bold = True

    c_right_bot = header_table.cell(1, 1)
    p_tn = c_right_bot.paragraphs[0]
    p_tn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tn.paragraph_format.space_after = Pt(0)
    r_tn = p_tn.add_run("Độc lập - Tự do - Hạnh phúc\n───────────────────────")
    r_tn.font.size = Pt(12)
    r_tn.font.bold = True

    # Spacer
    sp = document.add_paragraph()
    sp.paragraph_format.space_before = Pt(30)
    sp.paragraph_format.space_after = Pt(10)

    # Tiêu đề chính trang bìa
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("E-HỒ SƠ MỜI THẦU")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 32, 96)  # Navy công vụ

    p_sub = document.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    r_sub = p_sub.add_run("(Áp dụng hình thức Đấu thầu rộng rãi qua mạng — Một giai đoạn một túi hồ sơ)")
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True

    # Khung bảng tóm tắt thông tin gói thầu trên trang bìa
    cover_box = document.add_table(rows=4, cols=2)
    cover_box.style = "Table Grid"
    cover_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_box.autofit = False

    box_widths = [Mm(45), Mm(115)]
    for row in cover_box.rows:
        for i, w in enumerate(box_widths):
            row.cells[i].width = w
            _set_cell_margins(row.cells[i], top=100, bottom=100, left=150, right=150)

    labels = [
        ("Tên gói thầu:", doc.package.package_name),
        ("Chủ đầu tư:", doc.package.investor or "Theo Quyết định phê duyệt KHLCNT"),
        ("Mã định danh hồ sơ:", doc.doc_id),
        ("Căn cứ pháp lý:", "Luật Đấu thầu 22/2023/QH15, Nghị định 214/2025/NĐ-CP, TT 22/2024/TT-BKHĐT"),
    ]

    for idx, (label, val) in enumerate(labels):
        cell_lbl = cover_box.cell(idx, 0)
        cell_val = cover_box.cell(idx, 1)

        p1 = cell_lbl.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(12)

        p2 = cell_val.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(str(val))
        r2.font.size = Pt(12)
        if idx == 0:
            r2.font.bold = True

    # Cảnh báo dự thảo
    p_warn = document.add_paragraph()
    p_warn.paragraph_format.space_before = Pt(25)
    p_warn.paragraph_format.space_after = Pt(15)
    p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_w = p_warn.add_run(
        "⚠️ LƯU Ý: Dự thảo hồ sơ do Hệ thống AI Trợ lý Đấu thầu AutoTender-VN hỗ trợ tạo lập.\n"
        "Bắt buộc thực hiện thẩm định và phê duyệt theo đúng thẩm quyền trước khi phát hành chính thức."
    )
    r_w.font.size = Pt(11)
    r_w.font.italic = True
    r_w.font.bold = True
    r_w.font.color.rgb = RGBColor(180, 0, 0)

    # Chân trang bìa
    p_date = document.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(20)
    current_year = datetime.now().year
    current_month = datetime.now().month
    r_d = p_date.add_run(f"NĂM {current_year}")
    r_d.font.bold = True
    r_d.font.size = Pt(13)

    document.add_page_break()

    # =========================================================================
    # PHẦN II: CẢNH BÁO NẾU CÒN MỤC CHƯA DUYỆT
    # =========================================================================
    approved_count, total_count = doc.approval_progress
    if approved_count < total_count:
        h_warn = document.add_heading("⚠️ CẢNH BÁO: TÀI LIỆU CÒN MỤC CHƯA PHÊ DUYỆT", level=1)
        h_warn.paragraph_format.space_before = Pt(12)
        h_warn.paragraph_format.space_after = Pt(6)

        p_stat = document.add_paragraph(f"Tiến độ: Đã phê duyệt {approved_count}/{total_count} mục. Các mục sau CHƯA được phê duyệt:")
        p_stat.paragraph_format.space_after = Pt(6)

        for s in doc.sections:
            if s.status != "approved":
                p_item = document.add_paragraph(f"• {s.section_id} — {s.title} (Trạng thái: {s.status})")
                p_item.paragraph_format.space_after = Pt(2)
                p_item.paragraph_format.left_indent = Mm(10)
        document.add_page_break()

    ordered_sections = _ordered_sections(doc.sections)

    # =========================================================================
    # PHẦN III: MỤC LỤC CHÍNH THỨC (TABLE OF CONTENTS)
    # =========================================================================
    h_toc = document.add_heading("MỤC LỤC HỒ SƠ MỜI THẦU", level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_toc.paragraph_format.space_before = Pt(12)
    h_toc.paragraph_format.space_after = Pt(12)

    chapters: dict[str, list] = {}
    for s in ordered_sections:
        chapter = s.section_id.split(".")[0]
        chapters.setdefault(chapter, []).append(s)

    for chapter, sections in chapters.items():
        p_ch = document.add_paragraph()
        p_ch.paragraph_format.space_before = Pt(6)
        p_ch.paragraph_format.space_after = Pt(2)
        r_ch = p_ch.add_run(f"• {CHAPTER_TITLES.get(chapter, chapter).upper()}")
        r_ch.font.bold = True
        r_ch.font.size = Pt(12)

        for s in sections:
            p_sec = document.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(1)
            p_sec.paragraph_format.space_after = Pt(2)
            p_sec.paragraph_format.left_indent = Mm(10)
            r_sec = p_sec.add_run(f"- {s.section_id}: {s.title}")
            r_sec.font.size = Pt(11.5)

    document.add_page_break()

    # =========================================================================
    # PHẦN IV: NỘI DUNG TOÀN VĂN 8 CHƯƠNG I - VIII (CHUẨN THỂ THỨC NĐ 30/2020)
    # =========================================================================
    for chapter, sections in chapters.items():
        # Tiêu đề Chương (Heading 1)
        h_c = document.add_heading(CHAPTER_TITLES.get(chapter, chapter).upper(), level=1)
        h_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_c.paragraph_format.space_before = Pt(14)
        h_c.paragraph_format.space_after = Pt(10)
        for r in h_c.runs:
            r.font.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0, 32, 96)

        for s in sections:
            # Tiêu đề Mục (Heading 2)
            h_s = document.add_heading(f"{s.section_id}. {s.title}", level=2)
            h_s.paragraph_format.space_before = Pt(10)
            h_s.paragraph_format.space_after = Pt(4)
            for r in h_s.runs:
                r.font.bold = True
                r.font.size = Pt(13)
                r.font.color.rgb = RGBColor(0, 0, 0)

            # Phân tách đoạn văn bản, canh đều hai bên (Justify), thụt đầu dòng 1.27cm
            raw_paragraphs = [p_text.strip() for p_text in s.current_text.split("\n") if p_text.strip()]
            for p_text in raw_paragraphs:
                p_body = document.add_paragraph()
                p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_body.paragraph_format.first_line_indent = Mm(12.7)  # Thụt lề 1.27cm theo NĐ 30/2020
                p_body.paragraph_format.line_spacing = 1.25          # Giãn dòng 1.25 line
                p_body.paragraph_format.space_after = Pt(6)          # Cách đoạn 6pt

                r_body = p_body.add_run(p_text)
                r_body.font.name = "Times New Roman"
                r_body.font.size = Pt(13)

        document.add_page_break()

    # =========================================================================
    # PHẦN V: PHỤ LỤC NHẬT KÝ THẨM ĐỊNH & PHÊ DUYỆT (AUDIT LOG APPENDIX)
    # =========================================================================
    h_app = document.add_heading("PHỤ LỤC — NHẬT KÝ THẨM ĐỊNH & PHÊ DUYỆT (HITL)", level=1)
    h_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_app.paragraph_format.space_before = Pt(14)
    h_app.paragraph_format.space_after = Pt(10)

    p_app_intro = document.add_paragraph(
        "Bảng tổng hợp lưu vết quá trình rà soát, chỉnh sửa và phê duyệt của các chuyên gia "
        "đối với từng điều khoản trong hồ sơ mời thầu:"
    )
    p_app_intro.paragraph_format.space_after = Pt(8)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    t_widths = [Mm(25), Mm(55), Mm(25), Mm(25), Mm(30)]
    for i, w in enumerate(t_widths):
        table.rows[0].cells[i].width = w

    # Header Bảng: Tô nền xám nhạt #EAEAEA, Đậm, Canh giữa
    hdr_titles = ["Mục", "Tiêu đề", "Trạng thái", "Người duyệt", "Thời điểm"]
    hdr_cells = table.rows[0].cells
    for i, name in enumerate(hdr_titles):
        _set_cell_background(hdr_cells[i], "EAEAEA")
        _set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(name)
        r.font.bold = True
        r.font.size = Pt(11)

    # Thêm dữ liệu từng dòng
    for row_data in store.get_approval_log(doc.doc_id):
        row_cells = table.add_row().cells
        for i, w in enumerate(t_widths):
            row_cells[i].width = w
            _set_cell_margins(row_cells[i], top=80, bottom=80, left=80, right=80)

        values = [
            (row_data.get("section_id", ""), WD_ALIGN_PARAGRAPH.CENTER),
            (row_data.get("title", ""), WD_ALIGN_PARAGRAPH.LEFT),
            (row_data.get("status", "").upper(), WD_ALIGN_PARAGRAPH.CENTER),
            (row_data.get("approved_by") or "-", WD_ALIGN_PARAGRAPH.CENTER),
            (row_data.get("approved_at") or "-", WD_ALIGN_PARAGRAPH.CENTER),
        ]

        for i, (val_text, align) in enumerate(values):
            p = row_cells[i].paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val_text))
            r.font.size = Pt(10.5)

    document.save(str(output_path))
    return output_path
