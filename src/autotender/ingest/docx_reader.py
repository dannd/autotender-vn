"""Trích xuất text từ DOCX (Mục 6/M1) — dùng `python-docx`."""

from __future__ import annotations

from pathlib import Path

from autotender.utils.vn_text import normalize_document_text


def extract_docx_text(docx_path: str | Path) -> str:
    """Trích toàn bộ text từ DOCX, giữ thứ tự đoạn văn và bảng (nếu có)."""
    import docx  # python-docx — import trễ để module import được kể cả khi chưa cài lib

    document = docx.Document(str(docx_path))
    parts: list[str] = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return normalize_document_text("\n".join(parts))
